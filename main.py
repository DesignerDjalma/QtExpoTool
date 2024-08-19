import os
import sys
import time
import webbrowser
from getpass import getuser
from pathlib import Path
from datetime import datetime
import comtypes.client # type: ignore
import pandas as pd # type: ignore
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QAction,
    QApplication,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMainWindow,
    QPushButton,
    QVBoxLayout,
    QWidget,
    QToolButton,
    QMenu,
    QTableView,
    QProgressBar,
)


def timer_decorator(func):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        elapsed_time = end_time - start_time
        print(
            f"[bold green]Tempo de execução: {elapsed_time:.2f} segundos[/bold green]"
        )
        return result

    return wrapper


def set_doc_margins(doc, top, right, bottom, left):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.right_margin = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)


class Worker(QThread):
    progress = pyqtSignal(int)

    def __init__(self, doc, df, chunk_size=1000):
        super().__init__()
        self.doc = doc
        self.df = df
        self.chunk_size = chunk_size

    def run(self):
        self.add_data_to_doc(self.doc, self.df, self.chunk_size)

    def add_data_to_doc(self, doc, df, chunk_size):
        if not isinstance(df, pd.DataFrame):
            return

        t_size = len(df)
        num_chunks = (t_size + chunk_size - 1) // chunk_size  # Número de chunks
        count = 0

        for i in range(num_chunks):
            count += 1
            start_row = i * chunk_size
            end_row = min((i + 1) * chunk_size, t_size)
            df_chunk = df.iloc[start_row:end_row]

            # Adicionar cabeçalho
            if i == 0:
                table = doc.tables[1]
                table.style = "EstiloPadrao"
                hdr_cells = table.rows[0].cells

                for j, column in enumerate(df.columns):
                    hdr_cells[j].text = column
                    xml_bg_cinza = parse_xml(
                        r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w"))
                    )
                    hdr_cells[j]._element.get_or_add_tcPr().append(xml_bg_cinza)
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.bold = True
                        paragraph.alignment = 1

            # Adicionar dados
            for idx, (_, row) in enumerate(df_chunk.iterrows()):
                t_rows = ((count - 1) * chunk_size) + idx
                row_cells = table.add_row().cells

                for j, value in enumerate(row):
                    row_cells[j].text = str(value)
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                        paragraph.alignment = 1

                # Atualizando a barra de progresso
                progress_percentage = ((t_rows + 1) / t_size) * 100
                self.progress.emit(int(progress_percentage))

            doc.add_paragraph("")  # Adicionar uma quebra de linha entre os lotes

        print("Processo concluído!")


class TableModel(QtCore.QAbstractTableModel):
    def __init__(self, data, headers):
        super(TableModel, self).__init__()
        self._data = data
        self._headers = headers

    def data(self, index, role):
        if role == Qt.DisplayRole:
            return self._data[index.row()][index.column()]

    def rowCount(self, index):
        return len(self._data)

    def columnCount(self, index):
        return len(self._data[0])

    def setHeaderData(self, section, orientation, value, role=Qt.DisplayRole):
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                self._headers[section] = value
                return True
        return False

    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return self._headers[section]

    @staticmethod
    def formatValuesDataFrame(df_column: pd.Series) -> pd.Series:
        return (
            df_column.str.replace(",", ".")
            .str.replace(" m", "")
            .astype(float)
            .round(2)
            .apply(lambda x: f"{x:.2f} m".replace(".", ","))
        )


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.df = None
        self.rootdir = Path(__file__).parent
        self.icons_folder = rf"{self.rootdir}\src\icons"

        self.setWindowTitle("ExpoTool")
        self.setGeometry(100, 100, 440, 540)

        self.setWindowIcon(QIcon(rf"{self.icons_folder}\datatransformation.png"))

        self.create_menu()
        self.create_toolbar()
        self.statusBar().showMessage("Pronto")
        self.create_central_widget()

    def create_menu(self):
        menubar = self.menuBar()

        file_menu = menubar.addMenu("Arquivo")
        about_menu = menubar.addMenu("Ajuda")

        import_xlsx_action  = QAction(QIcon(rf"{self.icons_folder}\select_xlsx.ico"), "Importar .xlsx", self)
        import_docxmodel_action = QAction(QIcon(rf"{self.icons_folder}\select_docx.ico"), "Importar documento Word",self,)
        export_xlsx_action  = QAction(QIcon(rf"{self.icons_folder}\select_folder.ico"), "Exportar .xlsx", self)
        export_docx_action  = QAction(QIcon(rf"{self.icons_folder}\xlsx.ico"), "Exportar .docx", self)
        export_pdf_action   = QAction(QIcon(rf"{self.icons_folder}\docx.ico"), "Exportar .pdf", self)
        exit_action         = QAction(QIcon(rf"{self.icons_folder}\pdf.ico"), "Sair", self)
        about_action        = QAction(QIcon(rf"{self.icons_folder}\about.ico"), "Sobre", self)
        github_action       = QAction(QIcon(rf"{self.icons_folder}\github.ico"), "Github", self)

        import_xlsx_action.triggered.connect(self.f_import_excel)
        import_docxmodel_action.triggered.connect(self.f_import_docxmodel)
        export_xlsx_action.triggered.connect(self.f_export_xlsx)
        export_docx_action.triggered.connect(self.f_export_docx)
        export_pdf_action.triggered.connect(self.f_export_pdf)
        about_action.triggered.connect(self.f_about)
        exit_action.triggered.connect(self.close)
        github_action.triggered.connect(self.f_github)

        file_menu.addAction(import_xlsx_action)
        file_menu.addAction(import_docxmodel_action)
        file_menu.addSeparator()
        file_menu.addAction(export_xlsx_action)
        file_menu.addAction(export_docx_action)
        file_menu.addAction(export_pdf_action)
        file_menu.addSeparator()
        file_menu.addAction(exit_action)
        about_menu.addAction(about_action)
        about_menu.addAction(github_action)

    def save_doc(self, doc, docx_output):
        doc.save(docx_output)
        QtWidgets.QMessageBox.information(
            self,
            "Exportação",
            Rf"Dados exportados com sucesso para: {docx_output}",
        )
        self.worker.progress.emit(int(0))

    def f_github(self):
        url = "https://github.com/DesignerDjalma/QtExpoTool"
        sim = QtWidgets.QMessageBox.question(
                self,
                "Abrindo Repositório",
                Rf"Você está prestar a ir para o reposotório Github no seu navegador padrão. Deseja continuar?",
            )
        nao = QtWidgets.QMessageBox.question(
                self,
                "Abrindo Repositório",
                Rf"Você está prestar a ir para o reposotório Github no seu navegador padrão. Deseja continuar?",
            )
        print(sim)
        print(nao)
        x = 1
        webbrowser.open(url)

    def f_export_docx(self):

        if not self.export_path.text():
            QtWidgets.QMessageBox.warning(
                self,
                "Local da exportação",
                Rf"Nenhum pasta de exportação foi selecionada! Selecione uma pasta de exportação.",
            )
            return
        if not isinstance(self.df, pd.DataFrame):
            return

        if not self.df.empty:
            docx_output = rf"{self.export_path.text()}/DOCUMENTO_{datetime.now().strftime('%H%M%S')}.docx"
            docx = Document(self.docxmodel_path.text())

            # Configura o Worker e a barra de progresso
            self.worker = Worker(docx, self.df)
            self.worker.progress.connect(self.updateProgress)

            # Inicia o processamento
            self.worker.start()

            # Aguarda a conclusão do processamento
            self.worker.finished.connect(lambda: self.save_doc(docx, docx_output))
        else:
            QtWidgets.QMessageBox.warning(self, "Aviso", "Nenhum dado para exportar!")

    def f_export_xlsx(self):

        if not self.export_path.text():
            QtWidgets.QMessageBox.warning(
                self,
                "Local da exportação",
                Rf"Nenhum pasta de exportação foi selecionada! Selecione uma pasta de exportação.",
            )
            return
        if not isinstance(self.df, pd.DataFrame):
            return
        df_export = self.df.copy()
        xlsx_output = rf"{self.export_path.text()}/PLANILHA_{datetime.now().strftime('%H%M%S')}.xlsx"
        df_export.to_excel(xlsx_output, index=False)
        QtWidgets.QMessageBox.information(
            self,
            "Exportação",
            f"Dados exportados com sucesso para: {xlsx_output}",
        )

    def f_export_pdf(self):
        if not self.export_path.text():
            QtWidgets.QMessageBox.warning(
                self,
                "Local da exportação",
                Rf"Nenhum pasta de exportação foi selecionada! Selecione uma pasta de exportação.",
            )
            return
        if not isinstance(self.df, pd.DataFrame):
            return
        pdf_output = (
            rf"{self.export_path.text()}/PDF_{datetime.now().strftime('%H%M%S')}.pdf"
        )
        QtWidgets.QMessageBox.critical(
            self,
            "Exportação",
            Rf"Função não implementada ainda, em breve em atualizações.",
        )

    def f_about(self):
        QtWidgets.QMessageBox.information(
            self,
            "Sobre",
            f"Olá {getuser().capitalize()}, Obrigado por utilizar o QtExpoTool 2024 v1.0.0",
        )

    def create_toolbar(self):
        toolbar                 = self.addToolBar("Ferramentas")

        import_xlsx_action      = QAction(QIcon(rf"{self.icons_folder}\select_xlsx.ico"),"Importar arquivo Excel",self,)
        import_docxmodel_action = QAction(QIcon(rf"{self.icons_folder}\select_docx.ico"),"Importar documento Word",self,)
        export_folder_action    = QAction(QIcon(rf"{self.icons_folder}\select_folder.ico"),"Selecionar Pasta de Exportação",self,)
        export_xlsx             = QAction(QIcon(rf"{self.icons_folder}\xlsx.ico"), "Exportar Excel (XLSX)", self)
        export_docx             = QAction(QIcon(rf"{self.icons_folder}\docx.ico"), "Exportar (DOCX)", self)
        export_pdf              = QAction(QIcon(rf"{self.icons_folder}\pdf.ico"), "Exportar (PDF)", self)
        about_action            = QAction(QIcon(rf"{self.icons_folder}\about.ico"), "Sobre", self)
        github_action           = QAction(QIcon(rf"{self.icons_folder}\github.ico"), "Abrir Repositório Github", self)
        
        import_xlsx_action.triggered.connect(self.f_import_excel)
        import_docxmodel_action.triggered.connect(self.f_import_docxmodel)
        export_folder_action.triggered.connect(self.f_export_location)
        export_xlsx.triggered.connect(self.f_export_xlsx)
        export_docx.triggered.connect(self.f_export_docx)
        export_pdf.triggered.connect(self.f_export_pdf)
        about_action.triggered.connect(self.f_about)
        github_action.triggered.connect(self.f_github)

        toolbar.addAction(import_xlsx_action)
        toolbar.addAction(import_docxmodel_action)
        toolbar.addAction(export_folder_action)
        toolbar.addSeparator()
        toolbar.addAction(export_xlsx)
        toolbar.addAction(export_docx)
        toolbar.addAction(export_pdf)
        toolbar.addSeparator()
        toolbar.addAction(about_action)
        toolbar.addAction(github_action)

    def create_central_widget(self):
        self.table_initial_data = [["Aguardando", "Dados do", "Usuário"]]
        self.table = QTableView()
        self.model = TableModel(
            self.table_initial_data, ["Coluna A", "Coluna B", "Coluna C"]
        )
        self.table.setModel(self.model)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.my_layout = QVBoxLayout()

        self.btn_export_all = QPushButton("Exportar")
        self.btn_export_all.setMaximumWidth(100)
        self.btn_export_all.setLayoutDirection(1)
        self.btn_export_all (self.export_all)

        # Section: Import .xlsx
        self.import_docxmodel_layout = QHBoxLayout()
        self.docxmodel_label = QLabel("Relatorio Modelo:")
        self.docxmodel_label.setMinimumWidth(100)
        self.docxmodel_path = QLineEdit()
        self.docxmodel_path.setReadOnly(True)

        self.docxmodel_tool_button = QToolButton()
        self.docxmodel_tool_button.setText("... ")

        self.menu_docxmodel_tool_button = QMenu()
        self.select_docxmodel_action = self.menu_docxmodel_tool_button.addAction(
            "Selecionar Documento"
        )

        self.select_docxmodel_action.triggered.connect(self.f_import_docxmodel)
        self.docxmodel_tool_button.setMenu(self.menu_docxmodel_tool_button)
        self.docxmodel_tool_button.setPopupMode(QToolButton.InstantPopup)

        self.import_docxmodel_layout.addWidget(self.docxmodel_label)
        self.import_docxmodel_layout.addWidget(self.docxmodel_path)
        self.import_docxmodel_layout.addWidget(self.docxmodel_tool_button)

        # Section: Import .xlsx
        self.import_excel_layout = QHBoxLayout()
        self.excel_label = QLabel("Planilha Excel:")
        self.excel_label.setMinimumWidth(100)
        self.excel_path = QLineEdit()
        self.excel_path.setReadOnly(True)

        self.excel_tool_button = QToolButton()
        self.excel_tool_button.setText("... ")

        self.menu_excel_tool_button = QMenu()
        self.select_excel_action = self.menu_excel_tool_button.addAction(
            "Selecionar Planilha"
        )

        self.select_excel_action.triggered.connect(self.f_import_excel)
        self.excel_tool_button.setMenu(self.menu_excel_tool_button)
        self.excel_tool_button.setPopupMode(QToolButton.InstantPopup)

        self.import_excel_layout.addWidget(self.excel_label)
        self.import_excel_layout.addWidget(self.excel_path)
        self.import_excel_layout.addWidget(self.excel_tool_button)

        # Section: Export path
        self.export_path_layout = QHBoxLayout()
        self.export_label = QLabel("Local de Exportação:")
        self.export_label.setMinimumWidth(100)
        self.export_path = QLineEdit()
        self.export_path.setReadOnly(True)

        self.export_tool_button = QToolButton()
        self.export_tool_button.setText("... ")

        self.menu_export_tool_button = QMenu()
        self.select_local_action = self.menu_export_tool_button.addAction(
            "Localizar Local Exportação"
        )

        self.select_local_action.triggered.connect(self.f_export_location)
        self.export_tool_button.setMenu(self.menu_export_tool_button)
        self.export_tool_button.setPopupMode(QToolButton.InstantPopup)

        self.export_path_layout.addWidget(self.export_label)
        self.export_path_layout.addWidget(self.export_path)
        self.export_path_layout.addWidget(self.export_tool_button)

        self.progressBarPercentage = QProgressBar(self)
        self.progressBarPercentage.setValue(0)

        self.my_layout.addLayout(self.import_docxmodel_layout)
        self.my_layout.addLayout(self.import_excel_layout)
        self.my_layout.addLayout(self.export_path_layout)
        self.my_layout.addWidget(self.table)
        self.my_layout.addWidget(self.progressBarPercentage)
        self.my_layout.addWidget(self.btn_export_all)

        self.central_widget.setLayout(self.my_layout)

    def f_import_docxmodel(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Selecione o Documento",
            "",
            "Document Files (*.docx; *.doc)",
            options=options,
        )
        if file_path:
            self.docxmodel_path.setText(file_path)

    def export_all(self):
        self.statusBar().showMessage("Botão Exportar clicado!")

    def f_import_excel(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Selecione a Planilha Excel",
            "",
            "Excel Files (*.xlsx; *.xls)",
            options=options,
        )
        if file_path:
            self.excel_path.setText(file_path)
            self.df = pd.read_excel(file_path)
            self.df = self.df.astype(str)

            numeros = []
            self.distancias_sum_listmode = self.df["Distância"].tolist()
            for dist in self.distancias_sum_listmode:
                print(f"{dist = }")
                numeros.append(round(float(f"{dist}".replace(" m", "").replace(",", ".")), 2))


            self.distancia_sum = (
                self.df["Distância"]
                .str.replace(",", ".")
                .str.replace(" m", "")
                .astype(float)
                .round(2)
                .sum()
            )
            ps = f"{self.distancia_sum:.2f}".replace(".", ",")
            print(f"Soma do perímetro: {ps}")

            self.updateTableView(self.df)
            self.df["Distância"] = TableModel.formatValuesDataFrame(
                self.df["Distância"]
            )
            perimetro_sum_after = (
                self.df["Distância"]
                .str.replace(",", ".")
                .str.replace(" m", "")
                .astype(float)
                .round(2)
                .sum()
            )
            ps_after = f"{perimetro_sum_after:.2f}".replace(".", ",")
            self.statusBar().showMessage(
                f"Soma do perímetro: {ps} m | Total de linhas: {self.df.shape[0]} | Soma do perímetro após arrendondamento: {sum(numeros)}"
            )
    # Função para transformar o valor em float com duas casas decimais
    def transform_value(value):
        value = value.replace(' m', '').replace(',', '.')
        return round(float(value), 2)

    def updateTableView(self, df: pd.DataFrame):
        headers = df.columns.tolist()
        if len(df) >= 11:
            data_top = df.head()
            data_mid = ["..." for i in range(len(headers))]
            data_bot = df.tail()
            data = data_top.values.tolist() + [data_mid] + data_bot.values.tolist()
        else:
            data = df.values.tolist()

        self.model = TableModel(data, headers)
        self.table.setModel(self.model)

        for col in range(len(headers)):
            self.model.setHeaderData(col, Qt.Horizontal, headers[col])  # type: ignore

        self.model.layoutChanged.emit()
        self.table.setColumnWidth(7, 123)  # Coluna Latitude  ajustada
        self.table.setColumnWidth(8, 123)  # Coluna Longitude ajustada
        self.table.update()

    def f_export_location(self):
        options = QFileDialog.Options()
        folder_path = QFileDialog.getExistingDirectory(
            self, "Selecione o Local de Exportação", options=options
        )
        if folder_path:
            self.export_path.setText(folder_path)

    def updateProgress(self, value):
        """Update de Progress Bar % value, async"""
        self.progressBarPercentage.setValue(value)




if __name__ == "__main__":

    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
